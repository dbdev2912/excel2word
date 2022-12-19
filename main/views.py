from django.shortcuts import render
from django.http import JsonResponse, FileResponse
from django.http.response import StreamingHttpResponse
# Create your views here.
import json
import os

import docx
import openpyxl
import qrcode

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import io



def mainpage(request, *argv, **kwargs):
    return render(request, "home.html")


def generating_file(request, *argv, **kwargs):
    if request.method == "POST":
        sheet = request.POST.get("sheet")
        sheet = json.loads(sheet)

        data = sheet.get("data")
        valueField = sheet.get("valueField")
        labelField = sheet.get("labelField")
        aspect = sheet.get("aspect")
        with_index = sheet.get("with_index")
        export_file = sheet.get("export_name")
        if not export_file:
            export_file = "python-expoted-file.docx"
        try:
            os.remove(export_file)
        except:
            pass

        document = Document()
        x = aspect["x"]
        y = aspect["y"]
        cell_width = float(aspect["inches"])

        print(aspect["inches"])

        rows = 1 + (len(data) // x);
        if len(data) % x == 0:
            rows -= 1
        table = document.add_table(  rows, x, style="Table Grid")

        index = 0

        for i in range( rows ):
            row = table.rows[i].cells
            try:
                for j in range( 3 ):
                    cell_value = data[index][valueField]
                    cell_label = data[index][labelField]
                    qr_img = qrcode.make(cell_value)
                    qr_img = qr_img.crop((15, 25, 360, 350))
                    imgdata = io.BytesIO()
                    qr_img.save(imgdata, format="png")

                    para = row[j].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()

                    font  = run.font
                    font.bold = True
                    font.size = Pt(12)
                    if with_index:
                        cell_label = "{0}.{1}".format(index + 1, cell_label)

                    run.add_picture(imgdata, width=Inches(cell_width), height=Inches(cell_width))
                    run.add_text(cell_label)

                    index += 1
                    print("Inserted the {1}th image with value: {0}".format(cell_value, index))
            except:
                print("ERROR somewhere")

        document.save('main/export/{0}'.format(export_file))

        context = {
            "path": export_file,
            "success": True,
        }
        return JsonResponse(context)
    else:

        context = {
            "path": None,
            "success": False,
        }
        return JsonResponse(context)


def download_and_remove_file(request, file="", *argv, **kwargs):

    if request.method == "GET":
        export_file = file
        docx = open('main/export/{0}'.format(file), 'rb')
        response = FileResponse(
            docx,  # use the stream's content
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment;filename={0}'.format(export_file)
        return response
